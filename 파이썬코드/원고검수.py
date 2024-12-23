try:
    from docx import Document
    from docx.shared import RGBColor
    import gspread
    from oauth2client.service_account import ServiceAccountCredentials
    import os
    from pathlib import Path
    from datetime import datetime
    import win32com.client as win32
    import winreg
except ImportError as e:
    print(f"필요한 라이브러리를 설치해주세요: {e}")
    print("pip install lxml==4.9.3")
    print("pip install python-docx")
    print("pip install gspread oauth2client")
    print("pip install pywin32")
    exit(1)

def setup_hwp_security():
    """한글 보안 모듈 설정"""
    try:
        winup_path = r"Software\HNC\HwpAutomation\Modules"
        winreg.CreateKey(winreg.HKEY_CURRENT_USER, winup_path)
        key_path = winreg.OpenKey(winreg.HKEY_CURRENT_USER, winup_path, 0, winreg.KEY_ALL_ACCESS)
        winreg.SetValueEx(key_path, "FilePathCheckDLL", 0, winreg.REG_SZ, "")
        winreg.CloseKey(key_path)
    except Exception as e:
        print(f"보안 설정 중 오류: {str(e)}")

def highlight_keywords_hwp(hwp_path, keyword_notes, output_path):
    """한글 파일 처리"""
    try:
        # 한글 실행
        hwp = win32.gencache.EnsureDispatch("HWPFrame.HwpObject")
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        
        # 한글 파일 열기
        hwp.Open(os.path.abspath(hwp_path))
        
        # 각 키워드에 대해 처리
        for keyword, note in keyword_notes.items():
            # 찾아바꾸기 설정
            find_ctrl = hwp.HParameterSet.HFindReplace
            hwp.HAction.GetDefault("AllReplace", find_ctrl.HSet)
            
            # 키워드 찾기 설정
            find_ctrl.FindString = keyword
            find_ctrl.ReplaceString = keyword
            find_ctrl.IgnoreCase = 1
            
            # 글자 속성 설정 (빨간색, 굵게)
            find_ctrl.ReplaceCharShape.TextColor = hwp.RGBColor(251, 65, 65)
            find_ctrl.ReplaceCharShape.Bold = 1
            
            # 찾아바꾸기 실행
            hwp.HAction.Execute("AllReplace", find_ctrl.HSet)
            
            # 노트가 있는 경우 추가
            if note:
                find_ctrl.FindString = keyword
                find_ctrl.ReplaceString = f"{keyword} {note}"
                find_ctrl.ReplaceCharShape.TextColor = hwp.RGBColor(92, 179, 56)
                find_ctrl.ReplaceCharShape.Bold = 0
                hwp.HAction.Execute("AllReplace", find_ctrl.HSet)
        
        # 저장 및 종료
        hwp.SaveAs(os.path.abspath(output_path))
        hwp.Quit()
        print("한글 문서 처리가 완료되었습니다.")
        
    except Exception as e:
        print(f"한글 파일 처리 중 오류: {str(e)}")
        if 'hwp' in locals():
            hwp.Quit()

def convert_txt_to_docx(txt_path):
    """txt 파일을 docx로 변환"""
    try:
        # txt 파일 읽기 (여러 인코딩 시도)
        text = None
        for encoding in ['utf-8', 'cp949', 'euc-kr']:
            try:
                with open(txt_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            print(f"파일 인코딩을 확인할 수 없습니다: {txt_path}")
            return None
            
        # docx 파일 생성
        doc = Document()
        doc.add_paragraph(text)
        
        # 임시 docx 파일 경로
        docx_path = txt_path.replace('.txt', '_temp.docx')
        doc.save(docx_path)
        
        return docx_path
        
    except Exception as e:
        print(f"변환 중 오류 발생: {str(e)}")
        return None

def highlight_keywords(doc_path, keyword_notes, output_path):
    """파일 형식에 따라 적절한 처리 함수 호출"""
    try:
        # 파일 존재 확인
        if not os.path.exists(doc_path):
            print(f"Error: 파일을 찾을 수 없습니다 - {doc_path}")
            return
            
        # 파일 확장자 확인
        file_ext = os.path.splitext(doc_path)[1].lower()
        
        if file_ext == '.hwp':
            print(f"한글 파일은 처리하지 않습니다: {doc_path}")
            return
        elif file_ext == '.docx':
            # docx 파일 처리
            doc = Document(doc_path)
            
            # 모든 단락을 순회
            for paragraph in doc.paragraphs:
                # 단락의 텍스트 저장
                text = paragraph.text
                
                # 키워드 위치 찾기
                positions = []
                for keyword in keyword_notes.keys():
                    start = 0
                    while True:
                        index = text.find(keyword, start)
                        if index == -1:
                            break
                        positions.append((index, index + len(keyword), keyword))
                        start = index + 1
                
                if positions:
                    # 위치를 정렬
                    positions.sort()
                    
                    # 기존 runs 제거
                    for run in paragraph.runs:
                        run._element.getparent().remove(run._element)
                    
                    # 새로운 runs 추가
                    current_pos = 0
                    for start, end, keyword in positions:
                        # 키워드 전 텍스트
                        if start > current_pos:
                            run = paragraph.add_run(text[current_pos:start])
                        
                        # 키워드 (빨간색으로, 굵게)
                        run = paragraph.add_run(keyword)
                        run.font.color.rgb = RGBColor(251, 65, 65)
                        run.bold = True  # 굵게 표시
                        
                        # 키워드 옆에 노트 추가
                        if keyword_notes[keyword]:  # 노트가 있는 경우에만
                            note_run = paragraph.add_run(f" {keyword_notes[keyword]}")
                            note_run.font.color.rgb = RGBColor(92, 179, 56)
                        
                        current_pos = end
                    
                    # 마지��� 키워드 이후 텍스트
                    if current_pos < len(text):
                        run = paragraph.add_run(text[current_pos:])
            
            # 수정된 문서 저장
            doc.save(output_path)
            print("Word 문서 처리가 완료되었습니다.")
        elif file_ext == '.txt':
            print("txt 파일을 docx로 변환 중...")
            docx_path = convert_txt_to_docx(doc_path)
            
            if docx_path:
                # 변환된 docx 파일 처리
                docx_output = output_path.replace('.txt', '.docx')
                result = highlight_keywords(docx_path, keyword_notes, docx_output)
                
                # 임시 파일 삭제
                try:
                    os.remove(docx_path)
                except:
                    pass
                    
                return result
            else:
                print("txt 파일 변환 실패")
                return False
        else:
            print(f"지원하지 않는 파일 형식입니다: {file_ext}")
            
    except Exception as e:
        print(f"오류 발생: {str(e)}")

def get_keywords_from_sheet():
    """구글 시트에서 키워드와 사유를 가져오는 함수"""
    try:
        # 구글 시트 인증
        scope = ['https://spreadsheets.google.com/feeds',
                'https://www.googleapis.com/auth/drive']
        creds = ServiceAccountCredentials.from_json_keyfile_name(
            'D:/이채윤 파일/코딩/colab-408723-89110ae33a5b.json', 
            scope
        )
        client = gspread.authorize(creds)
        
        # 시트 열기
        sheet = client.open_by_url(
            'https://docs.google.com/spreadsheets/d/1eNCbstSMyQAA7CPvwb2qE7kZWg40B7Jf-fJ7ti0ABOE/edit?gid=0'
        ).worksheet('키워드')
        
        # 키워드와 사유/제안 가져오기
        keywords = sheet.col_values(2)[2:]  # B3부터
        reasons = sheet.col_values(3)[2:]   # C3부터
        
        # 딕셔너리로 변환
        keyword_notes = {}
        for keyword, reason in zip(keywords, reasons):
            if keyword.strip():  # 빈 셀 제외
                keyword_notes[keyword] = reason if reason else ''
                
        return keyword_notes
        
    except Exception as e:
        print(f"구글 시트 데이터 가져오기 실패: {str(e)}")
        return None

def find_file_with_extension(base_path):
    """파일 확장자 자동 찾기"""
    # 지원하는 확장자 목록 (.hwp 제외)
    extensions = ['.docx', '.txt']
    
    # 확장자가 없는 경로에 각 확장자를 붙여서 시도
    for ext in extensions:
        test_path = f"{base_path}{ext}"
        if os.path.exists(test_path):
            return test_path, ext
            
    # 파일을 찾지 못한 경우
    return None, None

if __name__ == "__main__":
    try:
        # 구글 시트에서 키워드와 사유 가져오기
        keyword_notes = get_keywords_from_sheet()
        if not keyword_notes:
            print("키워드를 가져오지 못했습니다.")
            exit(1)
            
        # 구글 시트 연결
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_name(
            'D:/이채윤 파일/코딩/colab-408723-89110ae33a5b.json', 
            scope
        )
        client = gspread.authorize(creds)
        
        # 검수파일 시트 열기
        sheet = client.open_by_url(
            'https://docs.google.com/spreadsheets/d/1eNCbstSMyQAA7CPvwb2qE7kZWg40B7Jf-fJ7ti0ABOE/edit?gid=226778372'
        ).worksheet('검수파일')
        
        # 파일 정보 가져오기
        file_paths = sheet.range('F4:F100')  # 충분히 큰 범위 지정
        file_names = sheet.range('G4:G100')
        output_paths = sheet.range('H4:H100')
        
        # 각 파일 처리
        for i, (path_cell, name_cell, output_cell) in enumerate(zip(file_paths, file_names, output_paths)):
            if path_cell.value and name_cell.value:  # 값이 있는 행만 처리
                base_path = f"{path_cell.value}\{name_cell.value}"
                input_file, ext = find_file_with_extension(base_path)
                
                if input_file and os.path.exists(input_file):
                    # 출력 파일에도 같은 확장자 사용
                    output_file = f"{path_cell.value}\{output_cell.value}{ext}"
                    
                    print(f"\n처리 중: {name_cell.value}{ext}")
                    highlight_keywords(input_file, keyword_notes, output_file)
                    
                    # 업데이트 일자 기록
                    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    sheet.update_cell(i + 4, 9, now)  # I열��� 업데이트 일자 기록
                    print(f"완료: {name_cell.value}{ext}")
                else:
                    print(f"\n파일을 찾을 수 없음: {base_path}")
                    print("지원하는 확장자: .txt, .docx")
            elif not path_cell.value:  # 빈 행을 만나면 종료
                break
                
        print("\n모든 파일 처리 완료")
        
    except Exception as e:
        print(f"오류 발생: {str(e)}")